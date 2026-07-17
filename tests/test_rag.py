"""Test RAG: chia đoạn, đọc file, và tìm kiếm FTS5 — tất cả offline."""

from io import BytesIO

import pytest

import db
from rag import chunk_text, extract_text


# ── Chia đoạn (chunking) ──────────────────────────────────────────────────
def test_van_ban_ngan_thanh_1_chunk():
    assert chunk_text("Xin chào") == ["Xin chào"]


def test_gop_cac_doan_van_ngan():
    text = "Đoạn một.\n\nĐoạn hai.\n\nĐoạn ba."
    chunks = chunk_text(text, max_chars=800)
    assert len(chunks) == 1  # cả 3 đoạn ngắn gộp thành 1 chunk
    assert "Đoạn một." in chunks[0] and "Đoạn ba." in chunks[0]


def test_khong_chunk_nao_qua_dai():
    text = "\n\n".join(f"Đoạn văn số {i} " + "nội dung " * 30 for i in range(20))
    chunks = chunk_text(text, max_chars=800)
    assert all(len(c) <= 800 for c in chunks)
    assert len(chunks) > 1


def test_doan_van_sieu_dai_bi_cat_co_overlap():
    text = "x" * 2000  # 1 đoạn văn liền 2000 ký tự, không có chỗ ngắt
    chunks = chunk_text(text, max_chars=800, overlap=100)
    assert all(len(c) <= 800 for c in chunks)
    # Tổng ký tự > 2000 vì có phần chồng lấn giữa các chunk
    assert sum(len(c) for c in chunks) > 2000


def test_van_ban_rong():
    assert chunk_text("") == []
    assert chunk_text("   \n\n  ") == []


# ── Đọc file ──────────────────────────────────────────────────────────────
def test_doc_file_txt():
    data = "Nội dung tiếng Việt".encode("utf-8")
    assert extract_text("ghi_chu.txt", data) == "Nội dung tiếng Việt"


def test_doc_file_docx():
    # Tạo file Word thật trong RAM bằng python-docx rồi đọc lại
    from docx import Document

    document = Document()
    document.add_paragraph("Điều 1: Nghỉ phép năm 12 ngày.")
    document.add_paragraph("Điều 2: Làm việc từ xa tối đa 2 ngày mỗi tuần.")
    buffer = BytesIO()
    document.save(buffer)

    text = extract_text("quy_dinh.docx", buffer.getvalue())
    assert "Nghỉ phép năm 12 ngày" in text
    assert "Làm việc từ xa" in text


def test_dinh_dang_khong_ho_tro():
    with pytest.raises(ValueError):
        extract_text("anh.jpg", b"...")


# ── Tìm kiếm FTS5 (db) ───────────────────────────────────────────────────
def test_tim_kiem_tai_lieu():
    db.add_document(111, "quy_dinh.docx", [
        "Điều 1: Nhân viên được nghỉ phép năm 12 ngày, chưa kể lễ tết.",
        "Điều 2: Giờ làm việc từ 8h30 đến 17h30, nghỉ trưa 1 tiếng.",
        "Điều 3: Làm việc từ xa tối đa 2 ngày mỗi tuần, đăng ký trước.",
    ])

    results = db.search_chunks(111, "nghỉ phép bao nhiêu ngày")
    assert results, "phải tìm thấy kết quả"
    content, doc_name = results[0]
    assert "12 ngày" in content  # đoạn liên quan nhất đứng đầu (BM25)
    assert doc_name == "quy_dinh.docx"


def test_tim_khong_dau_van_thay():
    db.add_document(111, "a.txt", ["Chính sách nghỉ phép năm của công ty."])
    assert db.search_chunks(111, "nghi phep") != []  # gõ không dấu


def test_khong_tim_thay_tai_lieu_nguoi_khac():
    db.add_document(999, "b.txt", ["Tài liệu bí mật về nghỉ phép của người khác."])
    assert db.search_chunks(111, "nghỉ phép") == []


def test_cau_hoi_rong_khong_crash():
    assert db.search_chunks(111, "") == []
    assert db.search_chunks(111, "!!! ???") == []


def test_liet_ke_va_xoa_tai_lieu():
    doc_id = db.add_document(111, "a.txt", ["đoạn 1", "đoạn 2"])
    khac = db.add_document(999, "cua_nguoi_khac.txt", ["x"])

    docs = db.list_documents(111)
    assert len(docs) == 1
    assert docs[0][1] == "a.txt" and docs[0][2] == 2  # tên + số đoạn

    assert db.delete_document(111, khac) is False  # không xóa được của người khác
    assert db.delete_document(111, doc_id) is True
    assert db.list_documents(111) == []
    assert db.search_chunks(111, "đoạn") == []  # chunks cũng bị xóa theo
